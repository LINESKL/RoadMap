import json
from docx import Document
import openai
import uuid
from django.conf import settings
import datetime

# Настройка API-ключа OpenAI (замените на свой ключ)
openai.api_key = settings.OPENAI_API_KEY
def extract_text_from_docx(file_path: str) -> str:
    """
    Извлекает весь текст из DOCX-файла, включая абзацы и таблицы.
    """
    try:
        doc = Document(file_path)
        full_text = []

        # Извлекаем текст из абзацев
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = '\n'.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        full_text.append(cell_text)

        return '\n'.join(full_text)

    except Exception as e:
        print(f"Ошибка при извлечении текста: {str(e)}")
        raise

def structure_data_with_gpt(text: str) -> dict:
    """
    Отправляет текст в GPT и просит создать структуру с предметом, 15 неделями и подтемами, генерируя задания.
    """
    prompt = f"""
    Ты помощник, который преобразует текст силлабуса в структурированную базу данных в формате JSON с объектом "roadmap". Текст силлабуса описывает курс, который длится 15 недель, с тематическим планом, включающим главные темы и подтемы. Твоя задача:

    1. Создай объект "roadmap" с тремя вложенными полями:
       - "subject": объект, представляющий предмет.
         - "label": название предмета (например, "Философия").
         - "description": краткое описание курса из раздела "Краткое описание курса".
       - "weeks": список объектов, представляющих 15 недель курса.
         - "id": строка вида "week-1", "week-2" и т.д. (ровно 15 записей).
         - "week_number": номер недели (от 1 до 15).
         - "main_topic": главная тема недели (из тематического плана).
         - "description": краткое описание или результаты обучения для недели.
         - "next_week_id": id следующей недели (например, "week-2" для "week-1", или null для "week-15").
       - "subtopics": список объектов, представляющих подтемы внутри недель.
         - "id": строка вида "subtopic-1", "subtopic-2" и т.д.
         - "week_id": ссылка на id недели (например, "week-1").
         - "label": название подтемы (например, "Обзор литературы").
         - "description": описание подтемы или результаты обучения.
         - "material": литература или материалы, связанные с подтемой (из списка литературы).
         - "assignment": задание, которое ты сгенерируешь самостоятельно на основе темы недели ("main_topic") и названия подтемы ("label"). Не бери задания из текста силлабуса, а придумай их сам, чтобы они были логично связаны с темой и подтемой. Задания должны быть конкретными, выполнимыми и соответствовать академическому контексту (например, написать эссе, подготовить презентацию, провести анализ и т.д.).
    2. Учти, что тематический план может содержать диапазоны недель (например, "1-3"). Раздели такие диапазоны на отдельные недели (1, 2, 3) с общей темой, но с возможными разными подтемами, если они указаны.
    3. Установи последовательную связь между неделями через "next_week_id" (week-1 -> week-2 -> ... -> week-15).
    4. Если данных недостаточно, заполни недостающие недели (до 15) минимальными значениями, используя общую тему курса.
    5. Возвращай ТОЛЬКО валидный JSON без дополнительных пояснений или форматирования (например, без "```json").

    Вот текст силлабуса:

    {text}
    """

    # Отправка запроса в GPT
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Ты помощник, который преобразует текст силлабуса в структурированную базу данных в формате JSON с объектом 'roadmap'. Возвращай ТОЛЬКО валидный JSON без пояснений."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4000,
        temperature=0.1
    )

    # Извлечение результата
    result = response.choices[0].message['content'].strip()

    # Попытка извлечь JSON
    try:
        json_start = result.find('{')
        json_end = result.rfind('}') + 1
        if json_start != -1 and json_end != -1:
            result = result[json_start:json_end]
        structured_data = json.loads(result)

        # Проверка на обязательные поля
        if 'roadmap' not in structured_data or 'subject' not in structured_data['roadmap'] or 'weeks' not in structured_data['roadmap'] or 'subtopics' not in structured_data['roadmap']:
            raise ValueError("GPT не вернул обязательный объект 'roadmap' с полями 'subject', 'weeks' или 'subtopics'")
        if len(structured_data['roadmap']['weeks']) != 15:
            raise ValueError(f"Ожидается ровно 15 недель, получено {len(structured_data['roadmap']['weeks'])}")
        return structured_data
    except json.JSONDecodeError as e:
        print(f"Ошибка декодирования JSON из ответа GPT: {e}")
        print(f"Полученный текст: {result}")
        raise

def save_to_mongodb(structured_data: dict, db, file_path: str):
    """
    Сохраняет структурированные данные в MongoDB, обновляя существующий предмет или создавая новый.
    """
    syllabus_collection = db['syllabus']

    # Извлекаем название предмета из структуры
    subject_label = structured_data['roadmap']['subject']['label']

    # Проверяем, существует ли уже запись для этого предмета
    existing_syllabus = syllabus_collection.find_one({'roadmap.subject.label': subject_label})
    syllabus_id = str(uuid.uuid4()) if not existing_syllabus else existing_syllabus['id']

    # Подготавливаем данные для сохранения
    syllabus_data = {
        'id': syllabus_id,
        'file_path': file_path,
        'upload_date': datetime.datetime.now(),
        'processed': True,
        'roadmap': structured_data['roadmap']
    }

    # Обновляем или вставляем запись
    if existing_syllabus:
        syllabus_collection.update_one({'id': syllabus_id}, {'$set': syllabus_data})
        print(f"Обновлены данные для предмета: {subject_label}")
    else:
        syllabus_collection.insert_one(syllabus_data)
        print(f"Сохранены данные для нового предмета: {subject_label}")
